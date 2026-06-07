"""
Converts the Markdown report from the M&A agent into a .docx file
using python-docx (already in requirements.txt).
"""
import re
from io import BytesIO

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DELIVERABLE_LABELS = {
    "complet":   "Rapport Complet",
    "fiche":     "Fiche Entreprise",
    "benchmark": "Benchmark Concurrents",
    "manda":     "Note M&A & Secteur",
    "geo":       "Analyse Géographique",
}

NAVY   = RGBColor(0x1A, 0x27, 0x44)
BLUE   = RGBColor(0x2D, 0x5A, 0x7A)
DARK   = RGBColor(0x1A, 0x20, 0x2C)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT  = RGBColor(0xEE, 0xF2, 0xF6)


# ── Inline markdown parser ────────────────────────────────────────────────────

def _parse_inline(paragraph, text: str):
    """Add runs for **bold**, *italic*, `code` — handles nested edge cases."""
    # Improved pattern: allows any char except newline inside markers
    pattern = r'(\*\*(?:[^*]|\*(?!\*))+?\*\*|\*(?:[^*])+?\*|`[^`\n]+?`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**') and len(part) > 4:
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`') and len(part) > 2:
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Courier New'
            run.font.size = Pt(9)
        else:
            if part:
                paragraph.add_run(part)


def _add_border(paragraph, color='C8D4DC'):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _shade_paragraph(paragraph, fill='E8EDF2'):
    pPr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill)
    pPr.append(shd)


def _shade_cell(cell, fill='1A2744'):
    """Apply background colour to a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)


def _is_separator_row(line: str) -> bool:
    """True for Markdown table separator rows like |---|---| or |:---|:---:|."""
    return bool(re.match(r'^\|[\s\-:|]+\|[\s\-:|]*$', line.strip()))


def _add_md_table(doc, table_lines: list):
    """Render a list of Markdown table lines as a Word table."""
    # Separate header from data (skip separator rows)
    rows = []
    for line in table_lines:
        if _is_separator_row(line):
            continue
        cells = [c.strip() for c in line.strip().strip('|').split('|')]
        rows.append(cells)

    if not rows:
        return

    n_cols = max(len(r) for r in rows)
    tbl    = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = 'Table Grid'

    for r_idx, row_cells in enumerate(rows):
        for c_idx in range(n_cols):
            cell_text = row_cells[c_idx] if c_idx < len(row_cells) else ''
            cell = tbl.rows[r_idx].cells[c_idx]
            cell.text = ''
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)

            if r_idx == 0:
                # Header row — navy background, white bold text
                _shade_cell(cell, '1A2744')
                run = para.add_run(cell_text)
                run.bold = True
                run.font.color.rgb = WHITE
                run.font.size      = Pt(9)
            else:
                # Data row — alternate shading
                if r_idx % 2 == 0:
                    _shade_cell(cell, 'EEF2F6')
                _parse_inline(para, cell_text)
                for run in para.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()  # spacer after table


# ── Main generator ────────────────────────────────────────────────────────────

def generate_word(markdown_text: str, company_name: str, deliverable_type: str) -> bytes:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # ── Cover page ────────────────────────────────────────────────────────────
    label = DELIVERABLE_LABELS.get(deliverable_type, "Rapport M&A")

    cover_title = doc.add_heading('', level=0)
    cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cover_title.add_run("Screening M&A")
    run.font.color.rgb = NAVY
    run.font.size      = Pt(28)

    sub     = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(f"{company_name}  ·  {label}")
    sub_run.font.size      = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x4A, 0x7F, 0xA5)
    sub_run.bold = True

    doc.add_paragraph()  # spacer

    # ── Pre-clean: remove ━━━ separator lines outside code blocks ─────────────
    cleaned_lines = []
    in_pre = False
    for line in markdown_text.split('\n'):
        if line.strip().startswith('```'):
            in_pre = not in_pre
            cleaned_lines.append(line)
            continue
        if not in_pre and re.match(r'^━+$', line.strip()):
            cleaned_lines.append('---')  # replace with proper HR
        else:
            cleaned_lines.append(line)
    lines = cleaned_lines

    # ── Parse markdown body ───────────────────────────────────────────────────
    in_code_block = False
    code_lines: list = []

    i = 0
    while i < len(lines):
        raw = lines[i]

        # ── Code fence ──
        if raw.strip().startswith('```'):
            if not in_code_block:
                in_code_block = True
                code_lines    = []
            else:
                in_code_block = False
                if code_lines:
                    p = doc.add_paragraph('\n'.join(code_lines))
                    p.style = 'No Spacing'
                    for r in p.runs:
                        r.font.name  = 'Courier New'
                        r.font.size  = Pt(8.5)
                        r.font.color.rgb = NAVY
                    _shade_paragraph(p)
            i += 1
            continue

        if in_code_block:
            code_lines.append(raw)
            i += 1
            continue

        # ── Markdown table ──
        if raw.startswith('|') and '|' in raw[1:]:
            table_lines = [raw]
            j = i + 1
            while j < len(lines) and lines[j].startswith('|'):
                table_lines.append(lines[j])
                j += 1
            _add_md_table(doc, table_lines)
            i = j   # jump past all collected table lines
            continue

        # ── Headings ──
        if raw.startswith('# '):
            p = doc.add_heading(raw[2:].strip(), level=1)
            for r in p.runs:
                r.font.color.rgb = NAVY

        elif raw.startswith('## '):
            p = doc.add_heading(raw[3:].strip(), level=2)
            for r in p.runs:
                r.font.color.rgb = NAVY

        elif raw.startswith('### '):
            p = doc.add_heading(raw[4:].strip(), level=3)
            for r in p.runs:
                r.font.color.rgb = BLUE

        # ── Horizontal rule ──
        elif raw.strip() in ('---', '***', '___'):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _add_border(p)

        # ── Blockquote ──
        elif raw.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            _parse_inline(p, raw[2:])
            for r in p.runs:
                r.font.color.rgb = RGBColor(0x3A, 0x50, 0x68)
                r.italic = True

        # ── Unordered list ──
        elif re.match(r'^[-*•] ', raw):
            p = doc.add_paragraph(style='List Bullet')
            _parse_inline(p, re.sub(r'^[-*•] ', '', raw))

        # ── Ordered list ──
        elif re.match(r'^\d+\. ', raw):
            p = doc.add_paragraph(style='List Number')
            _parse_inline(p, re.sub(r'^\d+\. ', '', raw))

        # ── Empty line ──
        elif raw.strip() == '':
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

        # ── Normal paragraph ──
        else:
            p = doc.add_paragraph()
            _parse_inline(p, raw)

        i += 1

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
